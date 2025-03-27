import streamlit as st
import pandas as pd
import os
import tempfile
from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import networkx as nx
import matplotlib.pyplot as plt
from pyvis.network import Network
import streamlit.components.v1 as components
import json
import numpy as np
from langchain.schema import Document
import uuid
import pickle
import re
from transformers import pipeline

# Page configuration
st.set_page_config(page_title="Document Graph RAG", layout="wide", page_icon="📚")

# Initialize session state variables
if 'documents' not in st.session_state:
    st.session_state.documents = []
if 'knowledge_graph' not in st.session_state:
    st.session_state.knowledge_graph = nx.DiGraph()
if 'vector_store' not in st.session_state:
    st.session_state.vector_store = None
if 'knowledge_graph_built' not in st.session_state:
    st.session_state.knowledge_graph_built = False
if 'graph_html' not in st.session_state:
    st.session_state.graph_html = None
if 'ner_pipeline' not in st.session_state:
    st.session_state.ner_pipeline = None

# Load NER pipeline
@st.cache_resource
def load_ner_pipeline():
    return pipeline("ner", model="dslim/bert-base-NER")

# Initialize NLP components on first run
with st.spinner("Loading NLP models (first run only)..."):
    if st.session_state.ner_pipeline is None:
        st.session_state.ner_pipeline = load_ner_pipeline()

# File handling functions
def extract_text_from_file(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file.name.split(".")[-1]}') as tmp_file:
        tmp_file.write(file.getbuffer())
        file_path = tmp_file.name
    
    text = ""
    try:
        file_extension = file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            pdf_reader = PdfReader(file_path)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        elif file_extension == 'docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        elif file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        else:
            st.warning(f"Unsupported file type: {file_extension}. Only PDF, DOCX, and TXT files are supported.")
    
    except Exception as e:
        st.error(f"Error extracting text from {file.name}: {str(e)}")

    os.remove(file_path)
    
    return text

def process_documents(files):
    all_docs = []
    
    for file in files:
        text = extract_text_from_file(file)
        if text:
            metadata = {
                "source": file.name,
                "file_type": file.name.split('.')[-1],
                "file_size": file.size
            }
            
            all_docs.append(Document(page_content=text, metadata=metadata))
    
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    docs = []
    for doc in all_docs:
        chunks = text_splitter.split_documents([doc])
        docs.extend(chunks)
    
    return docs

# Entity extraction functions
def extract_entities_and_relationships(text):
    """Extract entities and infer relationships using rule-based approach and NER"""
    # Extract entities using NER
    try:
        # Process with NER pipeline
        ner_results = st.session_state.ner_pipeline(text)
        
        # Group entities
        current_entity = None
        current_type = None
        current_start = None
        entities = []
        
        for entity in ner_results:
            if entity['entity'].startswith('B-'):  # Beginning of entity
                if current_entity:  # Save previous entity if exists
                    entities.append({
                        'name': current_entity,
                        'label': current_type
                    })
                current_entity = entity['word'].replace('#', '')
                current_type = entity['entity'][2:]  # Remove B- prefix
                current_start = entity['start']
            elif entity['entity'].startswith('I-') and current_entity:  # Inside of entity
                # Only append if it's the same entity type
                if entity['entity'][2:] == current_type:
                    # Check if this token is right after the current entity
                    if entity['start'] >= current_start:
                        # Handle subword tokens that start with ##
                        if entity['word'].startswith('##'):
                            current_entity += entity['word'][2:]
                        else:
                            current_entity += ' ' + entity['word']
                        current_start = entity['end']
        
        # Add the last entity if there is one
        if current_entity:
            entities.append({
                'name': current_entity,
                'label': current_type
            })
        
        # Clean up entity names and standardize labels
        cleaned_entities = []
        for entity in entities:
            # Clean the entity name
            clean_name = entity['name'].strip()
            
            # Map NER labels to our expected format
            label_mapping = {
                'PER': 'Person',
                'LOC': 'Location',
                'ORG': 'Organization',
                'MISC': 'Concept'
            }
            
            label = label_mapping.get(entity['label'], entity['label'])
            
            if clean_name and len(clean_name) > 1:  # Only add non-empty entities
                cleaned_entities.append({
                    'name': clean_name,
                    'label': label
                })
        
        # Simple rule-based relationship extraction
        # Find entities that appear close to each other in text
        relationships = []
        
        # Only process if we have multiple entities
        if len(cleaned_entities) > 1:
            for i in range(len(cleaned_entities)):
                for j in range(i+1, len(cleaned_entities)):
                    entity1 = cleaned_entities[i]
                    entity2 = cleaned_entities[j]
                    
                    # Check if entities are close in the text (within 100 chars)
                    e1_pos = text.find(entity1['name'])
                    e2_pos = text.find(entity2['name'])
                    
                    if e1_pos >= 0 and e2_pos >= 0:
                        distance = abs(e1_pos - e2_pos)
                        
                        if distance < 100:
                            # Determine relationship type based on entity types
                            rel_type = "RELATED_TO"
                            
                            # Person -> Organization
                            if entity1['label'] == 'Person' and entity2['label'] == 'Organization':
                                rel_type = "WORKS_FOR"
                            # Organization -> Person
                            elif entity1['label'] == 'Organization' and entity2['label'] == 'Person':
                                rel_type = "EMPLOYS"
                            # Person -> Location
                            elif entity1['label'] == 'Person' and entity2['label'] == 'Location':
                                rel_type = "LOCATED_IN"
                            # Organization -> Location
                            elif entity1['label'] == 'Organization' and entity2['label'] == 'Location':
                                rel_type = "BASED_IN"
                            
                            # Add relationship
                            relationships.append({
                                'start': entity1['name'],
                                'type': rel_type,
                                'end': entity2['name']
                            })
        
        return {
            'entities': cleaned_entities,
            'relationships': relationships
        }
    
    except Exception as e:
        st.warning(f"Error in entity extraction: {str(e)}")
        return {'entities': [], 'relationships': []}

# Knowledge graph functions
def build_knowledge_graph(docs):
    try:
        # Initialize embeddings - use sentence transformers (free alternative)
        embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        
        # Create a vector store with document chunks
        vector_store = FAISS.from_documents(docs, embeddings)
        st.session_state.vector_store = vector_store
        
        # Initialize a new graph
        G = nx.DiGraph()
        
        # Create a mapping of chunks to IDs for reference
        chunk_to_id = {}
        
        # Process each document chunk and extract entities + relationships
        for i, doc in enumerate(docs):
            if i > 30:  # Limit for performance
                break
                
            # Generate a unique ID for the chunk
            chunk_id = f"chunk_{i}"
            chunk_to_id[doc.page_content[:50]] = chunk_id
            
            # Add the chunk to the graph
            G.add_node(chunk_id, 
                       type="Chunk", 
                       text=doc.page_content[:200], 
                       full_text=doc.page_content,
                       metadata=doc.metadata)
            
            # Use our local entity extractor
            data = extract_entities_and_relationships(doc.page_content)
            
            # Process entities
            entity_ids = {}  # Map entity names to their IDs
            
            for entity in data.get("entities", []):
                entity_label = entity['label']
                entity_name = entity['name']
                
                # Generate a unique ID for the entity (or reuse if it exists)
                entity_id = f"{entity_label}_{entity_name}"
                entity_ids[entity_name] = entity_id
                
                # Add or update entity in graph
                if not G.has_node(entity_id):
                    G.add_node(entity_id, type=entity_label, name=entity_name, label=entity_name)
                
                # Connect chunk to entity
                G.add_edge(chunk_id, entity_id, type="MENTIONS")
            
            # Process relationships
            for rel in data.get("relationships", []):
                start_name = rel['start']
                rel_type = rel['type']
                end_name = rel['end']
                
                # Check if both entities exist in our entity_ids
                if start_name in entity_ids and end_name in entity_ids:
                    start_id = entity_ids[start_name]
                    end_id = entity_ids[end_name]
                    
                    # Add the relationship to the graph
                    G.add_edge(start_id, end_id, type=rel_type)
        
        # Store the graph in session state
        st.session_state.knowledge_graph = G
        
        # Save graph to a temporary file for persistence
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pkl') as tmp_file:
            pickle.dump(G, tmp_file)
            st.session_state.graph_file = tmp_file.name
            
        return True
    
    except Exception as e:
        st.error(f"Error building knowledge graph: {str(e)}")
        return False

def visualize_graph():
    try:
        G = st.session_state.knowledge_graph
        
        # Check if graph is empty or only has chunk nodes
        entity_nodes = [n for n, attr in G.nodes(data=True) if attr.get('type') != 'Chunk']
        
        if not entity_nodes:
            st.warning("No entities found in the graph. Try adding more documents or adjusting the extraction process.")
            return None
        
        # Create network visualization
        net = Network(height="600px", width="100%", notebook=True, bgcolor="#222222", font_color="white")
        
        # Add entity nodes only
        for node_id in entity_nodes:
            node_data = G.nodes[node_id]
            node_type = node_data.get('type', 'Unknown')
            node_name = node_data.get('name', node_id)
            
            net.add_node(
                node_id, 
                label=node_name, 
                title=f"{node_type}: {node_name}",
                color=get_color_for_label(node_type)
            )
        
        # Add edges between entities
        edge_count = 0
        for source, target, data in G.edges(data=True):
            # Only include edges between entities (not chunks)
            if (source in entity_nodes and target in entity_nodes):
                edge_type = data.get('type', 'RELATED_TO')
                
                net.add_edge(
                    source, 
                    target, 
                    title=edge_type,
                    label=edge_type
                )
                edge_count += 1
        
        if edge_count == 0 and len(entity_nodes) > 0:
            st.info("Found entities but no relationships between them. The visualization will show disconnected entities.")
        
        # Set visualization options
        net.set_options("""
        {
          "physics": {
            "forceAtlas2Based": {
              "gravitationalConstant": -50,
              "centralGravity": 0.01,
              "springLength": 100,
              "springConstant": 0.08
            },
            "maxVelocity": 50,
            "solver": "forceAtlas2Based",
            "timestep": 0.35,
            "stabilization": {
              "enabled": true,
              "iterations": 1000
            }
          },
          "edges": {
            "smooth": {
              "type": "continuous",
              "forceDirection": "none"
            }
          },
          "interaction": {
            "navigationButtons": true,
            "keyboard": true
          }
        }
        """)
        
        # Save and read HTML
        path = tempfile.mkdtemp()
        net.save_graph(f"{path}/graph.html")
        
        with open(f"{path}/graph.html", "r", encoding="utf-8") as f:
            html = f.read()
        
        return html
        
    except Exception as e:
        st.error(f"Error visualizing graph: {str(e)}")
        return None

def get_color_for_label(label):
    label_colors = {
        "Person": "#3498db", #blue 
        "Organization": "#e74c3c", #red
        "Location": "#2ecc71", #green
        "Concept": "#9b59b6", #purple
        "Product": "#f39c12", #yellow
        "Event": "#1abc9c" #grene
    }
    return label_colors.get(label, "#95a5a6")

def answer_question(question):
    try:
        G = st.session_state.knowledge_graph
        vector_store = st.session_state.vector_store
        
        if not vector_store:
            return "No documents have been processed yet. Please upload and process documents first."
        
        # Retrieve relevant chunks
        search_results = vector_store.similarity_search(question, k=5)
        relevant_chunks = search_results
        
        # Find relevant entity nodes from the retrieved chunks
        relevant_entities = []
        chunk_texts = [chunk.page_content for chunk in relevant_chunks]
        
        # Get entities connected to these chunks
        for node, attr in G.nodes(data=True):
            if attr.get('type') == 'Chunk':
                chunk_text = attr.get('full_text', '')
                if any(chunk_text[:100] in full_chunk for full_chunk in chunk_texts):
                    # Get connected entities
                    for _, entity_id in G.out_edges(node):
                        entity_data = G.nodes[entity_id]
                        if entity_data.get('type') != 'Chunk':
                            entity_info = {
                                'type': entity_data.get('type'),
                                'name': entity_data.get('name')
                            }
                            if entity_info not in relevant_entities:
                                relevant_entities.append(entity_info)
        
        # Get relationships between relevant entities
        relationships = []
        entity_names = [e['name'] for e in relevant_entities]
        
        for source, target, data in G.edges(data=True):
            source_data = G.nodes.get(source, {})
            target_data = G.nodes.get(target, {})
            
            if (source_data.get('type') != 'Chunk' and 
                target_data.get('type') != 'Chunk' and
                source_data.get('name') in entity_names and
                target_data.get('name') in entity_names):
                
                relationships.append({
                    'start': source_data.get('name'),
                    'type': data.get('type'),
                    'end': target_data.get('name')
                })
        
        # Build context for answer generation
        context = "\n\n".join([chunk.page_content for chunk in relevant_chunks])
        
        # Create knowledge graph context
        kg_context = "Entities:\n"
        for entity in relevant_entities:
            kg_context += f"- {entity['type']}: {entity['name']}\n"
        
        kg_context += "\nRelationships:\n"
        for rel in relationships:
            kg_context += f"- {rel['start']} {rel['type']} {rel['end']}\n"
        
        # Simple rule-based answer generation (instead of OpenAI)
        answer = "Based on the documents, "
        
        # Find sentences in the context that contain keywords from the question
        question_keywords = set(question.lower().split()) - set(['a', 'an', 'the', 'is', 'are', 'was', 'were', 'be', 'been', 'being', 'to', 'of', 'and', 'in', 'that', 'have', 'with', 'for', 'this', 'from'])
        
        relevant_sentences = []
        for chunk in relevant_chunks:
            sentences = re.split(r'(?<=[.!?])\s+', chunk.page_content)
            for sentence in sentences:
                sentence_lower = sentence.lower()
                if any(keyword in sentence_lower for keyword in question_keywords):
                    relevant_sentences.append(sentence.strip())
        
        # Combine relevant sentences and entities into an answer
        if relevant_sentences:
            answer += " " + " ".join(relevant_sentences[:3])  # Use first 3 relevant sentences
        else:
            answer += " No direct information found in the documents."
        
        # Add entity information if available
        if relevant_entities:
            answer += "\n\nKey entities mentioned: "
            entity_info = []
            for entity in relevant_entities[:5]:  # Show top 5 entities
                entity_info.append(f"{entity['name']} ({entity['type']})")
            answer += ", ".join(entity_info)
        
        # Add relationship information if available
        if relationships:
            answer += "\n\nRelationships identified: "
            rel_info = []
            for rel in relationships[:3]:  # Show top 3 relationships
                rel_info.append(f"{rel['start']} {rel['type']} {rel['end']}")
            answer += "; ".join(rel_info)
        
        return answer
        
    except Exception as e:
        return f"Error answering question: {str(e)}"

# Main UI
st.title("📚 Knowledge Graph RAG System")

# File upload section
st.header("Step 1: Upload Documents")
uploaded_files = st.file_uploader("Upload documents related to your topic", 
                                 type=["pdf", "docx", "txt"], 
                                 accept_multiple_files=True)

if uploaded_files:
    if st.button("Process Documents"):
        with st.spinner("Processing documents..."):
            docs = process_documents(uploaded_files)
            st.session_state.documents = docs
            st.success(f"Processed {len(docs)} document chunks from {len(uploaded_files)} files.")

# Build knowledge graph section
st.header("Step 2: Build Knowledge Graph")
if st.session_state.documents:
    if st.button("Build Knowledge Graph"):
        with st.spinner("Building knowledge graph from documents... This may take a while."):
            success = build_knowledge_graph(st.session_state.documents)
            if success:
                st.session_state.knowledge_graph_built = True
                st.success("Knowledge graph successfully built!")
                # Generate visualization
                with st.spinner("Generating graph visualization..."):
                    st.session_state.graph_html = visualize_graph()
            else:
                st.error("Failed to build knowledge graph.")

# Explore knowledge graph section
st.header("Step 3: Explore Knowledge Graph")
if st.session_state.knowledge_graph_built:
    if st.session_state.graph_html:
        st.subheader("Knowledge Graph Visualization")
        # Create expander for the visualization
        with st.expander("Show Graph Visualization", expanded=True):
            components.html(st.session_state.graph_html, height=600)
            
        # Graph statistics
        G = st.session_state.knowledge_graph
        entity_count = len([n for n, attr in G.nodes(data=True) if attr.get('type') != 'Chunk'])
        relationship_count = len([e for e in G.edges() if G.nodes[e[0]].get('type') != 'Chunk' and G.nodes[e[1]].get('type') != 'Chunk'])
        
        col1, col2 = st.columns(2)
        with col1:
            st.metric("Entities", entity_count)
        with col2:
            st.metric("Relationships", relationship_count)
    else:
        st.warning("No visualization available. The graph might be empty or too complex to visualize.")

# Question answering section
st.header("Step 4: Ask Questions")
if st.session_state.knowledge_graph_built:
    question = st.text_input("Ask a question about your documents")
    if st.button("Get Answer") and question:
        with st.spinner("Searching the knowledge graph..."):
            answer = answer_question(question)
            st.markdown("### Answer:")
            st.write(answer)

# About section
with st.expander("About this application"):
    st.markdown("""
    ## Document Graph RAG System
    
    This application uses a knowledge graph-based RAG (Retrieval-Augmented Generation) approach to process and analyze documents. 
    
    ### How it works:
    
    1. **Document Processing**: Upload PDF, DOCX, or TXT files to extract and chunk the text.
    2. **Knowledge Graph Creation**: The system extracts entities and relationships from the documents using local models.
    3. **Graph Visualization**: Explore the entities and connections in your documents.
    4. **Question Answering**: Ask questions about your documents using the knowledge graph.
    
    ### Technologies:
    
    - **Streamlit**: For the web interface
    - **FAISS**: For vector embeddings and similarity search
    - **NetworkX**: For local graph storage and management
    - **HuggingFace Transformers**: For entity extraction using NER
    - **Sentence Transformers**: For text embeddings
    - **PyVis**: For graph visualization
    
    ### Usage Tips:
    
    - Upload related documents for better knowledge graph creation.
    - Keep questions specific to the content in your documents.
    - The knowledge graph visualization shows entities and relationships extracted from your documents.
    """)