import streamlit as st
import pandas as pd
import os
import tempfile
from PyPDF2 import PdfReader
import docx
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Neo4jVector
from langchain_community.graphs import Neo4jGraph
from langchain_community.llms import OpenAI
from langchain.chains import GraphQAChain
import networkx as nx
import matplotlib.pyplot as plt
from pyvis.network import Network
import streamlit.components.v1 as components
from neo4j import GraphDatabase
import json
import numpy as np
from langchain.schema import Document

st.set_page_config(page_title="Document Graph RAG", layout="wide", page_icon="📚")

if 'documents' not in st.session_state:
    st.session_state.documents = []
if 'knowledge_graph_built' not in st.session_state:
    st.session_state.knowledge_graph_built = False
if 'graph_html' not in st.session_state:
    st.session_state.graph_html = None

st.sidebar.title("Configuration")

with st.sidebar.expander("Neo4j Configuration", expanded=False):
    neo4j_uri = st.text_input("Neo4j URI", "bolt://localhost:7687")
    neo4j_username = st.text_input("Neo4j Username", "neo4j")
    neo4j_password = st.text_input("Neo4j Password", "Vociferous2003#", type="password")

with st.sidebar.expander("OpenAI Configuration", expanded=False):
    openai_api_key = st.text_input("OpenAI API Key ", "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 ",type="password")

def extract_text_from_file(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file.name.split(".")[-1]}') as tmp_file:
        tmp_file.write(file.getbuffer())
        file_path = tmp_file.name
    
    text = ""
    try:
        file_extension = file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            pdf_reader = PdfReader(file_path)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        elif file_extension == 'docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        elif file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        else:
            st.warning(f"Unsupported file type: {file_extension}. Only PDF, DOCX, and TXT files are supported.")
    
    except Exception as e:
        st.error(f"Error extracting text from {file.name}: {str(e)}")

    os.remove(file_path)
    
    return text

def process_documents(files):
    all_docs = []
    
    for file in files:
        text = extract_text_from_file(file)
        if text:
            metadata = {
                "source": file.name,
                "file_type": file.name.split('.')[-1],
                "file_size": file.size
            }
            
            all_docs.append(Document(page_content=text, metadata=metadata))
    

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    docs = []
    for doc in all_docs:
        chunks = text_splitter.split_documents([doc])
        docs.extend(chunks)
    
    return docs


def build_knowledge_graph(docs):
    try:
 
        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        

        driver = GraphDatabase.driver(neo4j_uri, auth=(neo4j_username, neo4j_password))
        with driver.session() as session:
            session.run("MATCH (n) DETACH DELETE n")
        driver.close()
        

        vectorstore = Neo4jVector.from_documents(
            docs,
            embeddings,
            url=neo4j_uri,
            username=neo4j_username,
            password=neo4j_password,
            index_name="document_chunks",
            node_label="Chunk",
            text_node_property="text",
            embedding_node_property="embedding"
        )
        llm = OpenAI(temperature=0, openai_api_key=openai_api_key)
        

        graph = Neo4jGraph(
            url=neo4j_uri,
            username=neo4j_username,
            password=neo4j_password
        )

        for i, doc in enumerate(docs):
            if i > 30:  
                break
                
            prompt = f"""
            Extract entities and relationships from the following text:
            
            {doc.page_content}
            
            Return the result as a JSON object with the following format:
            {{
                "entities": [
                    {{"label": "EntityType1", "name": "EntityName1"}},
                    {{"label": "EntityType2", "name": "EntityName2"}}
                ],
                "relationships": [
                    {{"start": "EntityName1", "type": "RELATIONSHIP_TYPE", "end": "EntityName2"}}
                ]
            }}
            
            Only include entity types like Person, Organization, Location, Concept, Product, etc.
            """
            
            try:
                response = llm.invoke(prompt)

                data = json.loads(response)

                entities_query = ""
                for entity in data.get("entities", []):
                    entities_query += f"""
                    MERGE (e:{entity['label']} {{name: "{entity['name']}"}})
                    """
                
                relationships_query = ""
                for rel in data.get("relationships", []):
                    relationships_query += f"""
                    MATCH (a {{name: "{rel['start']}"}})
                    MATCH (b {{name: "{rel['end']}"}})
                    MERGE (a)-[:{rel['type']}]->(b)
                    """

                if entities_query:
                    graph.query(entities_query)
                if relationships_query:
                    graph.query(relationships_query)

                chunk_id = i
                for entity in data.get("entities", []):
                    connect_query = f"""
                    MATCH (c:Chunk) WHERE id(c) = {chunk_id}
                    MATCH (e:{entity['label']} {{name: "{entity['name']}"}})
                    MERGE (c)-[:MENTIONS]->(e)
                    """
                    try:
                        graph.query(connect_query)
                    except:
                        pass  # Some chunks might not exist yet
            
            except Exception as e:
                st.warning(f"Error processing document {i}: {str(e)}")
                continue
        
        return True
    
    except Exception as e:
        st.error(f"Error building knowledge graph: {str(e)}")
        return False

def visualize_graph():
    try:
        graph = Neo4jGraph(
            url=neo4j_uri,
            username=neo4j_username,
            password=neo4j_password
        )

        nodes_query = """
        MATCH (n) 
        WHERE NOT n:Chunk
        RETURN id(n) as id, labels(n)[0] as label, n.name as name
        LIMIT 100
        """
        
        relationships_query = """
        MATCH (a)-[r]->(b)
        WHERE NOT a:Chunk AND NOT b:Chunk
        RETURN id(a) as source, id(b) as target, type(r) as type
        LIMIT 500
        """
        
        nodes_result = graph.query(nodes_query)
        relationships_result = graph.query(relationships_query)
        
        if not nodes_result or len(nodes_result) == 0:
            st.warning("No entities found in the graph. Try adding more documents or adjusting the extraction process.")
            return None

        net = Network(height="600px", width="100%", notebook=True, bgcolor="#222222", font_color="white")

        for node in nodes_result:
            net.add_node(
                node["id"], 
                label=node["name"], 
                title=f"{node['label']}: {node['name']}",
                color=get_color_for_label(node["label"])
            )
 
        for edge in relationships_result:
            net.add_edge(
                edge["source"], 
                edge["target"], 
                title=edge["type"],
                label=edge["type"]
            )

        net.set_options("""
        {
          "physics": {
            "forceAtlas2Based": {
              "gravitationalConstant": -50,
              "centralGravity": 0.01,
              "springLength": 100,
              "springConstant": 0.08
            },
            "maxVelocity": 50,
            "solver": "forceAtlas2Based",
            "timestep": 0.35,
            "stabilization": {
              "enabled": true,
              "iterations": 1000
            }
          },
          "edges": {
            "smooth": {
              "type": "continuous",
              "forceDirection": "none"
            }
          },
          "interaction": {
            "navigationButtons": true,
            "keyboard": true
          }
        }
        """)

        path = tempfile.mkdtemp()
        net.save_graph(f"{path}/graph.html")
        
        with open(f"{path}/graph.html", "r", encoding="utf-8") as f:
            html = f.read()
        
        return html
    
    except Exception as e:
        st.error(f"Error visualizing graph: {str(e)}")
        return None

def get_color_for_label(label):
    label_colors = {
        "Person": "#3498db",
        "Organization": "#e74c3c",
        "Location": "#2ecc71",
        "Concept": "#9b59b6",
        "Product": "#f39c12",
        "Event": "#1abc9c"
    }
    return label_colors.get(label, "#95a5a6")

def answer_question(question):
    try:
        graph = Neo4jGraph(
            url=neo4j_uri,
            username=neo4j_username,
            password=neo4j_password
        )
        
        llm = OpenAI(temperature=0, openai_api_key=openai_api_key)
        
        # Create GraphQAChain
        qa_chain = GraphQAChain.from_llm(
            llm=llm,
            graph=graph,
            verbose=True
        )
        
        # Get answer
        result = qa_chain.invoke({"query": question})
        
        return result["result"]
    
    except Exception as e:
        return f"Error answering question: {str(e)}"

st.title("📚 Document Graph RAG System")

# File upload section
st.header("Step 1: Upload Documents")
uploaded_files = st.file_uploader("Upload documents related to your topic", 
                                 type=["pdf", "docx", "txt"], 
                                 accept_multiple_files=True)

if uploaded_files:
    if st.button("Process Documents"):
        with st.spinner("Processing documents..."):
            docs = process_documents(uploaded_files)
            st.session_state.documents = docs
            st.success(f"Processed {len(docs)} document chunks from {len(uploaded_files)} files.")


st.header("Build Knowledge Graph")
if st.session_state.documents:
    if st.button("Build Knowledge Graph"):
        if not openai_api_key:
            st.error("Please enter your OpenAI API key in the sidebar.")
        else:
            with st.spinner("Building knowledge graph from documents... This may take a while."):
                success = build_knowledge_graph(st.session_state.documents)
                if success:
                    st.session_state.knowledge_graph_built = True
                    st.success("Knowledge graph successfully built!")
                    # Generate visualization
                    with st.spinner("Generating graph visualization..."):
                        st.session_state.graph_html = visualize_graph()
                else:
                    st.error("Failed to build knowledge graph.")

st.header("Explore Knowledge Graph")
if st.session_state.knowledge_graph_built:
    if st.session_state.graph_html:
        st.subheader("Knowledge Graph Visualization")
        # Create expander for the visualization
        with st.expander("Show Graph Visualization", expanded=True):
            components.html(st.session_state.graph_html, height=600)
    else:
        st.warning("No visualization available. The graph might be empty or too complex to visualize.")


st.header("Ask Questions")
if st.session_state.knowledge_graph_built:
    question = st.text_input("Ask a question about your documents")
    if st.button("Get Answer") and question:
        if not openai_api_key:
            st.error("Please enter your OpenAI API key in the sidebar.")
        else:
            with st.spinner("Searching the knowledge graph..."):
                answer = answer_question(question)
                st.markdown("### Answer:")
                st.write(answer)


with st.expander("About this application"):
    st.markdown("""
    ## Graph RAG (Retrieval Augmented Generation) System
    
    This application allows you to:
    1. Upload documents related to a specific topic
    2. Build a knowledge graph from these documents using Neo4j
    3. Visualize the entities and relationships extracted from your documents
    4. Ask questions and get answers based on the knowledge graph
    
    ### How it works:
    - Documents are processed and split into chunks
    - Each chunk is embedded and stored in Neo4j
    - Entities and relationships are extracted using language models
    - The knowledge graph is used to provide context-aware answers
    
    ### Requirements:
    - Neo4j database connection
    - OpenAI API key for embeddings and question answering
    
    ### Installation requirements:
    ```
    pip install streamlit pandas PyPDF2 python-docx langchain langchain-openai langchain-community neo4j networkx matplotlib pyvis
    ```
    """)