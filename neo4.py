import streamlit as st
import pandas as pd
import os
import tempfile
from PyPDF2 import PdfReader
import docx
from langchain_openai import OpenAIEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Neo4jVector
from langchain_community.graphs import Neo4jGraph
from langchain_community.llms import OpenAI
from langchain.chains import GraphQAChain
import networkx as nx
import matplotlib.pyplot as plt
from pyvis.network import Network
import streamlit.components.v1 as components
from neo4j import GraphDatabase
import json
import numpy as np
from langchain.schema import Document

st.set_page_config(page_title="Document Graph RAG", layout="wide", page_icon="📚")

if 'documents' not in st.session_state:
    st.session_state.documents = []
if 'knowledge_graph_built' not in st.session_state:
    st.session_state.knowledge_graph_built = False
if 'graph_html' not in st.session_state:
    st.session_state.graph_html = None

st.sidebar.title("Configuration")

with st.sidebar.expander("Neo4j Configuration", expanded=False):
    neo4j_uri = st.text_input("Neo4j URI", "bolt://localhost:7687")
    neo4j_username = st.text_input("Neo4j Username", "neo4j")
    neo4j_password = st.text_input("Neo4j Password", "Vociferous2003#", type="password")


with st.sidebar.expander("OpenAI Configuration", expanded=False):
    openai_api_key = st.text_input("OpenAI API Key ", "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 ",type="password")

# Function to initialize Neo4j driver
def get_neo4j_driver():
    try:
        driver = GraphDatabase.driver(neo4j_uri, auth=(neo4j_username, neo4j_password))
        driver.verify_connectivity()
        st.write("Neo4j connection successful!")
        return driver
    except Exception as e:
        st.error(f"Failed to connect to Neo4j: {str(e)}")
        return None
    
def extract_text_from_file(file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=f'.{file.name.split(".")[-1]}') as tmp_file:
        tmp_file.write(file.getbuffer())
        file_path = tmp_file.name
    
    text = ""
    try:
        file_extension = file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            pdf_reader = PdfReader(file_path)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        elif file_extension == 'docx':
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        
        elif file_extension == 'txt':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        else:
            st.warning(f"Unsupported file type: {file_extension}. Only PDF, DOCX, and TXT files are supported.")
    
    except Exception as e:
        st.error(f"Error extracting text from {file.name}: {str(e)}")

    os.remove(file_path)
    
    return text

def process_documents(files):
    all_docs = []
    
    for file in files:
        text = extract_text_from_file(file)
        if text:
            metadata = {
                "source": file.name,
                "file_type": file.name.split('.')[-1],
                "file_size": file.size
            }
            
            all_docs.append(Document(page_content=text, metadata=metadata))
    

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    
    docs = []
    for doc in all_docs:
        chunks = text_splitter.split_documents([doc])
        docs.extend(chunks)
    
    return docs
# Function to clear the Neo4j database
def clear_neo4j_database():
    try:
        driver = get_neo4j_driver()
        if driver:
            with driver.session() as session:
                session.run("MATCH (n) DETACH DELETE n")
            st.success("Neo4j database cleared successfully!")
        else:
            st.error("Unable to clear database: No connection to Neo4j.")
    except Exception as e:
        st.error(f"Error clearing Neo4j database: {str(e)}")

# Function to build the knowledge graph
def build_knowledge_graph(docs):
    try:
        st.write("Initializing Neo4j driver...")
        driver = get_neo4j_driver()
        if not driver:
            st.error("Failed to initialize Neo4j driver.")
            return False

        st.write("Clearing Neo4j database...")
        with driver.session() as session:
            session.run("MATCH (n) DETACH DELETE n")
        st.write("Database cleared successfully!")

        # Rest of the function...
    except Exception as e:
        st.error(f"Error in build_knowledge_graph: {str(e)}")
        return False

def visualize_graph():
    try:
        st.write("Initializing Neo4j driver for visualization...")
        driver = get_neo4j_driver()
        if not driver:
            st.error("Failed to initialize Neo4j driver.")
            return None

        st.write("Querying Neo4j for nodes and relationships...")
        with driver.session() as session:
            nodes_result = session.run("MATCH (n) RETURN id(n) as id, labels(n)[0] as label, n.name as name LIMIT 100").data()
            relationships_result = session.run("MATCH (a)-[r]->(b) RETURN id(a) as source, id(b) as target, type(r) as type LIMIT 500").data()

        st.write(f"Found {len(nodes_result)} nodes and {len(relationships_result)} relationships.")
        # Rest of the function...
    except Exception as e:
        st.error(f"Error in visualize_graph: {str(e)}")
        return None

def get_color_for_label(label):
    label_colors = {
        "Person": "#3498db",
        "Organization": "#e74c3c",
        "Location": "#2ecc71",
        "Concept": "#9b59b6",
        "Product": "#f39c12",
        "Event": "#1abc9c"
    }
    return label_colors.get(label, "#95a5a6")

def answer_question(question):
    try:
        driver = get_neo4j_driver()
        if not driver:
            st.error("Failed to initialize Neo4j driver.")
            return None

        llm = OpenAI(temperature=0, openai_api_key=openai_api_key)

        # Create GraphQAChain
        qa_chain = GraphQAChain.from_llm(
            llm=llm,
            graph=Neo4jGraph(
                url=neo4j_uri,
                username=neo4j_username,
                password=neo4j_password
            ),
            verbose=True
        )

        # Get answer
        result = qa_chain.invoke({"query": question})
        return result["result"]

    except Exception as e:
        return f"Error answering question: {str(e)}"

st.title("📚 Document Graph RAG System")

# File upload section
st.header("Step 1: Upload Documents")
uploaded_files = st.file_uploader("Upload documents related to your topic", 
                                 type=["pdf", "docx", "txt"], 
                                 accept_multiple_files=True)

if uploaded_files:
    if st.button("Process Documents"):
        with st.spinner("Processing documents..."):
            docs = process_documents(uploaded_files)
            st.session_state.documents = docs
            st.success(f"Processed {len(docs)} document chunks from {len(uploaded_files)} files.")


st.header("Build Knowledge Graph")
if st.session_state.documents:
    if st.button("Build Knowledge Graph"):
        if not openai_api_key:
            st.error("Please enter your OpenAI API key in the sidebar.")
        else:
            with st.spinner("Building knowledge graph from documents... This may take a while."):
                success = build_knowledge_graph(st.session_state.documents)
                if success:
                    st.session_state.knowledge_graph_built = True
                    st.success("Knowledge graph successfully built!")
                    # Generate visualization
                    with st.spinner("Generating graph visualization..."):
                        st.session_state.graph_html = visualize_graph()
                else:
                    st.error("Failed to build knowledge graph.")

st.header("Explore Knowledge Graph")
if st.session_state.knowledge_graph_built:
    if st.session_state.graph_html:
        st.subheader("Knowledge Graph Visualization")
        # Create expander for the visualization
        with st.expander("Show Graph Visualization", expanded=True):
            components.html(st.session_state.graph_html, height=600)
    else:
        st.warning("No visualization available. The graph might be empty or too complex to visualize.")


st.header("Ask Questions")
if st.session_state.knowledge_graph_built:
    question = st.text_input("Ask a question about your documents")
    if st.button("Get Answer") and question:
        if not openai_api_key:
            st.error("Please enter your OpenAI API key in the sidebar.")
        else:
            with st.spinner("Searching the knowledge graph..."):
                answer = answer_question(question)
                st.markdown("### Answer:")
                st.write(answer)


with st.expander("About this application"):
    st.markdown("""
    ## Graph RAG (Retrieval Augmented Generation) System
    
    This application allows you to:
    1. Upload documents related to a specific topic
    2. Build a knowledge graph from these documents using Neo4j
    3. Visualize the entities and relationships extracted from your documents
    4. Ask questions and get answers based on the knowledge graph
    
    ### How it works:
    - Documents are processed and split into chunks
    - Each chunk is embedded and stored in Neo4j
    - Entities and relationships are extracted using language models
    - The knowledge graph is used to provide context-aware answers
    
    ### Requirements:
    - Neo4j database connection
    - OpenAI API key for embeddings and question answering
    
    ### Installation requirements:
    ```
    pip install streamlit pandas PyPDF2 python-docx langchain langchain-openai langchain-community neo4j networkx matplotlib pyvis
    ```
    """)